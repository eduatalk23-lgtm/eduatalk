"""
TimeLevelUp 연구노트 생성기 (샘플 1건)
- 일자별 1페이지 docx
- git 로그 + session-handoff 요약 기반
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT_DIR = Path("/Users/johyeon-u/Desktop/coding/eduatalk/scripts/out/research-notes")
OUT_DIR.mkdir(parents=True, exist_ok=True)


def set_korean_font(run, size=10, bold=False):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    # 한글 폰트 적용 (East Asia)
    from docx.oxml.ns import qn
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "맑은 고딕")
    rFonts.set(qn("w:hAnsi"), "맑은 고딕")


def add_heading(doc, text, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=True)
    return p


def add_meta_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(f"{label}  ")
    set_korean_font(r1, size=10, bold=True)
    r2 = p.add_run(value)
    set_korean_font(r2, size=10)
    return p


def add_body(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_korean_font(run, size=size)
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.25
    run = p.runs[0] if p.runs else p.add_run()
    run.text = ""
    run = p.add_run(text)
    set_korean_font(run, size=size)
    return p


def build_note_2026_04_22():
    doc = Document()

    # 여백 최소화 → 1페이지 유지
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("연구노트")
    set_korean_font(run, size=18, bold=True)

    # 메타 정보
    add_meta_row(doc, "일자:", "2026년 4월 22일 (수)")
    add_meta_row(doc, "프로젝트:", "TimeLevelUp — AI 기반 학습 플랜·학종 컨설팅 플랫폼")
    add_meta_row(doc, "작성자:", "조현우")
    add_meta_row(doc, "연구 단계:", "G-6 트랙(Agent 보안 경계) 최종 Sprint")

    # 구분선 대체 빈 줄
    doc.add_paragraph()

    # 1. 연구 주제
    add_heading(doc, "1. 연구 주제", size=12)
    add_body(
        doc,
        "멀티테넌트 AI 에이전트 환경에서 superadmin(운영사 내부 관리자) 의 "
        "cross-tenant 접근 경계를 정의하고, 기존 tenant-scoped 보안 모델과의 양립 방안을 확정한다.",
    )

    # 2. 수행 내역
    add_heading(doc, "2. 수행 내역", size=12)
    add_bullet(doc, "Role Model 의사결정: superadmin = eduatalk 운영사 내부 관리자(Option A)로 확정. 고객사 최고 권한자는 admin(tenant-scoped)으로 분리.")
    add_bullet(doc, "사전 실사: A/B 혼재 지점 7곳 매핑 → 이미 대부분 A 정책으로 정렬됨을 확인하고 실 블로커 1건(resolveStudentTarget) 식별.")
    add_bullet(doc, "구현: resolveStudent.ts 에 superadmin 분기 신설(admin client 로 RLS 우회 후 학생 tenant_id 를 downstream 에 주입). /api/chat 의 컨텍스트 프롬프트에 cross-tenant scope 힌트 1줄 추가.")
    add_bullet(doc, "문서화: CLAUDE.md 에 'Multi-tenant Role Model' 섹션 신설(5 role × 스코프 × 대표 경로 매트릭스 + RLS 템플릿). roleFilter.ts 에 설계 의도 주석 8줄.")
    add_bullet(doc, "검증: 신규 테스트 6 + 회귀 1 케이스 추가(superadmin 매치 0/1/다수, tenant_id=null 방어, admin 회귀). lint 0 / build success / 5672 tests pass.")

    # 3. 주요 결과 및 발견
    add_heading(doc, "3. 주요 결과 · 발견", size=12)
    add_bullet(doc, "커밋 1건(3cad50f5) / 5 파일 / +281·-7 LOC / DB 마이그레이션 0건.")
    add_bullet(doc, "핵심 설계 결정: '학생 선택 자체가 tenant 주입의 자연 proxy' — superadmin 이 학생명을 호명하면 resolveStudent 가 해당 학생의 tenant 로 downstream 을 seed. 별도 tenant selector UI 불필요.")
    add_bullet(doc, "AgentContext.tenantId: string 필수 계약(Sprint 1) 유지 — superadmin 도 학생 문맥 진입 시 학생 tenant 로 seed 되므로 호환성 보존.")
    add_bullet(doc, "G-6 트랙 4개 Sprint(S1 DELETE RLS · S2 Subagent 감사 · S3 OTel tenant_id · S4 superadmin 경로) 전부 완결 → 트랙 ④(Agent 보안 경계) 종료.")

    # 4. 다음 계획
    add_heading(doc, "4. 다음 계획", size=12)
    add_bullet(doc, "(1순위) Phase D — 대화 압축 및 pgvector 기반 Memory 시스템 설계·MVP (1주+).")
    add_bullet(doc, "(2순위) α2 v2-pre 실측 — 김세린 + 인제고 1학년 교차 검증 (OpenAI 소액 예산 집행 후).")
    add_bullet(doc, "(후속) superadmin UX 폴리시 F1~F3 — searchStudentsAction 의 cross-tenant RPC 확장, artifact 자동 바인딩, RLS 정책 전수 감사(실수요 발생 시).")

    # 저장
    out_path = OUT_DIR / "2026-04-22.docx"
    doc.save(out_path)
    print(f"saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build_note_2026_04_22()
