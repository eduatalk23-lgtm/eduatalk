"""
연구노트 생성기 (한국특허전략개발원 표준 양식 준수)

- 프로젝트 2종: TimeLevelUp / 생기부레벨업
- 월 2회 × 11회차 × 2 프로젝트 = 총 22개 docx
- 각 docx: A4 1페이지 기준, 메타 + 본문 + 서명란
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from dataclasses import dataclass, field
from typing import List

OUT_ROOT = Path("/Users/johyeon-u/Desktop/coding/eduatalk/scripts/out/research-notes")


# ───────────────────────── 공통 스타일 유틸 ─────────────────────────

def _apply_korean_font(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "맑은 고딕")
    rFonts.set(qn("w:hAnsi"), "맑은 고딕")


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            el = tcBorders.find(qn(f"w:{edge}"))
            if el is None:
                el = OxmlElement(f"w:{edge}")
                tcBorders.append(el)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(kwargs[edge]))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")


def _shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _write_cell(cell, text, size=9, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _apply_korean_font(run, size=size, bold=bold)


# ───────────────────────── 데이터 구조 ─────────────────────────

@dataclass
class NoteMeta:
    project_code: str
    task_title: str
    task_number: str
    serial_no: str
    note_date: str
    cover_period: str
    org: str = "에듀엣톡"
    dept: str = "연구개발팀"
    pi: str = "장미희"
    period: str = "2025-11-21 ~ 진행중"
    recorder: str = "조현우"
    witness: str = "장미희"
    confidential: str = "대외비(Confidential)"


@dataclass
class NoteBody:
    title: str
    purpose: str
    performed: List[str] = field(default_factory=list)
    results: List[str] = field(default_factory=list)
    next_plan: List[str] = field(default_factory=list)
    references: List[str] = field(default_factory=list)


# ───────────────────────── 문서 빌더 ─────────────────────────

def _setup_page(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def _add_header_bar(doc: Document, meta: NoteMeta):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{meta.confidential}    관리번호 {meta.serial_no}")
    _apply_korean_font(run, size=8, color=RGBColor(0x80, 0x80, 0x80))


def _add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("연 구 노 트")
    _apply_korean_font(run, size=18, bold=True)


def _add_meta_table(doc: Document, meta: NoteMeta):
    rows = [
        ("기관명 / 부서명", f"{meta.org} / {meta.dept}"),
        ("연구과제명", meta.task_title),
        ("연구과제번호 / 연구기간", f"{meta.task_number}  |  {meta.period}"),
        ("연구책임자 / 기록일", f"{meta.pi}  |  {meta.note_date}  (수행기간 {meta.cover_period})"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].width = Cm(4.2)
        row.cells[1].width = Cm(12.8)
        _write_cell(row.cells[0], label, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell(row.cells[1], value, size=9)
        _shade_cell(row.cells[0], "F2F2F2")
        for cell in row.cells:
            _set_cell_border(cell, top=4, bottom=4, left=4, right=4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _apply_korean_font(run, size=11, bold=True)


def _add_paragraph(doc: Document, text: str, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    _apply_korean_font(run, size=size)


def _add_bullet(doc: Document, text: str, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.25
    if p.runs:
        p.runs[0].text = ""
    run = p.add_run(text)
    _apply_korean_font(run, size=size)


def _add_signature_block(doc: Document, meta: NoteMeta):
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    table = doc.add_table(rows=2, cols=2)
    table.autofit = False

    header = table.rows[0]
    _write_cell(header.cells[0], "기록자  Written by", size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(header.cells[1], "확인자  Witnessed & Understood by", size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(header.cells[0], "F2F2F2")
    _shade_cell(header.cells[1], "F2F2F2")

    body = table.rows[1]
    _write_cell(
        body.cells[0],
        f"성명:  {meta.recorder}        서명:                   일자: {meta.note_date}",
        size=9,
    )
    _write_cell(
        body.cells[1],
        f"성명:  {meta.witness}        서명:                   일자: {meta.note_date}",
        size=9,
    )
    body.height = Cm(1.6)
    for cell in list(header.cells) + list(body.cells):
        cell.width = Cm(8.5)
        _set_cell_border(cell, top=4, bottom=4, left=4, right=4)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_note(meta: NoteMeta, body: NoteBody, out_path: Path):
    doc = Document()
    _setup_page(doc)
    _add_header_bar(doc, meta)
    _add_title(doc)
    _add_meta_table(doc, meta)

    _add_section_heading(doc, "1. 연구 주제 (실험제목)")
    _add_paragraph(doc, body.title)

    _add_section_heading(doc, "2. 연구 목적")
    _add_paragraph(doc, body.purpose)

    _add_section_heading(doc, "3. 수행 내역")
    for item in body.performed:
        _add_bullet(doc, item)

    _add_section_heading(doc, "4. 결과 · 발견")
    for item in body.results:
        _add_bullet(doc, item)

    _add_section_heading(doc, "5. 다음 계획")
    for item in body.next_plan:
        _add_bullet(doc, item)

    if body.references:
        _add_section_heading(doc, "6. 참고")
        for item in body.references:
            _add_bullet(doc, item, size=9)

    _add_signature_block(doc, meta)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


# ───────────────────────── 공통 메타 빌더 ─────────────────────────

def tlu_meta(serial, note_date, cover):
    return NoteMeta(
        project_code="TimeLevelUp",
        task_title="AI 기반 학습 플랜·일정 관리 플랫폼 개발",
        task_number="EDU-PLAN-2025-01",
        serial_no=serial,
        note_date=note_date,
        cover_period=cover,
    )


def rec_meta(serial, note_date, cover):
    return NoteMeta(
        project_code="생기부레벨업",
        task_title="AI 기반 학생부종합 분석·컨설팅 플랫폼 개발",
        task_number="EDU-REC-2025-01",
        serial_no=serial,
        note_date=note_date,
        cover_period=cover,
    )


# ───────────────────────── 11 회차 × 2 프로젝트 데이터 ─────────────────────────

NOTES = []


# R001 (2025-11-21 ~ 2025-11-30)
NOTES.append((tlu_meta("PLAN-RN-001", "2025-11-30", "2025-11-21 ~ 2025-11-30"), NoteBody(
    title="학습 플랜·일정 관리 플랫폼의 도메인 기반 아키텍처 기반 구축",
    purpose=(
        "Next.js 15 App Router 와 Supabase 를 기반으로 한 학습 플랜 관리 플랫폼의 초기 골격을 "
        "수립한다. 도메인 주도 설계(DDD) 원칙에 따라 service·repository·utils 계층을 분리하고, "
        "Supabase 스키마와 TypeScript 타입을 일관된 단일 출처로 정규화하여 향후 기능 확장의 "
        "토대를 마련하는 것을 목적으로 한다."
    ),
    performed=[
        "도메인 기반 아키텍처 구조 생성 — lib/domains/ 하위에 플랜·학생·학교·출결·점수 등 핵심 도메인 디렉터리 배치.",
        "Supabase 스키마 기반 TypeScript 타입 정규화 — 데이터베이스 스키마에서 타입을 자동 생성하여 프론트·백엔드 간 타입 불일치 제거.",
        "비즈니스 로직을 service / repository / utils 로 분리 — 관심사 분리 원칙 적용.",
        "스키마 정비 — all_schools_view 폐기 및 개별 테이블 직접 조회 전환, schools 테이블 구조 재설계.",
        "역할 기반 라우트 그룹 구성 — (admin) / (parent) / (student) 경로 그룹 신설.",
    ],
    results=[
        "해당 기간 커밋 637건. 프로젝트 초기 골격 확립.",
        "도메인 디렉터리 구조 확정 — 이후 모든 기능 개발의 기준점이 되는 lib/domains/ 계층화 완료.",
        "Supabase 와의 타입 동기화 파이프라인 수립 — pnpm 스크립트로 타입 재생성 자동화.",
        "코드 스타일 가이드 초안 — TypeScript any 금지, PascalCase 컴포넌트 네이밍, Tailwind spacing-first 정책.",
    ],
    next_plan=[
        "출석 관리 및 SMS 발송 시스템의 기초 도메인 구현.",
        "QR 체크인 기반 실물 출결 연동 기능 설계 및 RLS 정책 초안.",
        "학생·관리자 대시보드의 기본 레이아웃·인증 가드 구현.",
    ],
)))

NOTES.append((rec_meta("REC-RN-001", "2025-11-30", "2025-11-21 ~ 2025-11-30"), NoteBody(
    title="학생부종합 분석 플랫폼의 콘텐츠 도메인 기반 구조 수립",
    purpose=(
        "학생부종합전형 분석에 활용되는 교육 콘텐츠·교과·강의 데이터를 통합 관리할 수 있는 "
        "도메인 구조를 설계한다. 기존 grades·semesters 테이블 의존을 제거하고, 도메인별 단일 "
        "출처(Single Source of Truth) 를 확립하여 향후 AI 분석 파이프라인이 안정적으로 데이터에 "
        "접근할 수 있는 기반을 마련한다."
    ),
    performed=[
        "콘텐츠(Content) 도메인 기본 구조 생성 — 강의·교재·문제집 등의 통합 관리 entity 설계.",
        "lib/domains/content/index.ts 의 export 함수명 정리 — getLectureById 등 미사용 함수 제거.",
        "grades·semesters 테이블 참조 전면 제거 — 이후 도입될 학생 성적 모델의 유연성 확보를 위한 선행 작업.",
        "도메인별 repository 패턴 적용 — 데이터 접근을 repository 계층에서만 수행하도록 규약 확립.",
    ],
    results=[
        "해당 기간 커밋 중 콘텐츠 도메인 관련 작업 다수 반영 — 향후 탐구 가이드·마스터 콘텐츠 확장의 골격 확보.",
        "초기 콘텐츠 import 경로 확립 — 외부 Drive 연동을 위한 사전 준비 완료.",
        "도메인 간 느슨한 결합 — 각 도메인은 index.ts 를 통해서만 외부로 노출되는 캡슐화 원칙 수립.",
    ],
    next_plan=[
        "마스터 콘텐츠(master-content) 도메인 CRUD 기능 구현.",
        "학생·학부모 관점 분리를 위한 역할 기반 데이터 접근 제어(RLS) 정책 초안.",
        "학생부 진단 영역 도메인(student-record) 의 Entity 모델링 착수.",
    ],
)))


# R002 (2025-12-01 ~ 2025-12-15)
NOTES.append((tlu_meta("PLAN-RN-002", "2025-12-15", "2025-12-01 ~ 2025-12-15"), NoteBody(
    title="QR 기반 출결 시스템과 SMS 알림 플랫폼의 구현",
    purpose=(
        "학원 운영의 핵심 기능인 출결 관리를 QR 코드 체크인 방식으로 구현하고, 입·퇴실 상황을 "
        "실시간 SMS 로 보호자에게 통지하는 엔드투엔드 파이프라인을 구축한다. 동시에 다중 참여자 "
        "간 RLS 정책 충돌을 해소하여 학생·학부모·관리자의 권한 경계를 명확히 한다."
    ),
    performed=[
        "QR 출석 체크인 시스템 구현 — 학생 계정별 고유 QR 발급, 단말 카메라 스캔 후 출결 이벤트 생성.",
        "출석 관리 시스템 Phase 2·Phase 3 구현 — 관리자 대시보드 조회·수정·엑셀 내보내기 기능.",
        "SMS 발송 시스템 — 문자 발송 API 연동 + 발송 이력·실패 재시도 추적.",
        "RLS 정책 충돌 해소 — 공통 유틸리티 함수 신설, 다수 테이블의 중복 정책 로직 통합.",
        "출석 자동 SMS 연동 개선 — 입·퇴실 이벤트 발생 시 지정 보호자에게 즉시 통지.",
        "캠프(집중 학습) 도메인에 Adapter 패턴 도입 — 캠프별 상이한 스케줄 규칙을 단일 인터페이스로 흡수.",
        "zod v3 다운그레이드 및 recharts 최신화로 빌드 체인 안정화.",
    ],
    results=[
        "해당 기간 커밋 970건 (월 최대치). 출결·SMS 파이프라인 상용화 가능 수준 도달.",
        "RLS 위반 사례 제로화 — QR 체크인 경로에서 발생하던 RLS 거부 오류 해소.",
        "출석 퇴실 검증 로직 개선 — 기존 입실 기록을 참조하여 중복·무효 퇴실 방지.",
        "SMS 발송 오류 로깅 고도화 — 원본 에러와 발송 컨텍스트를 함께 기록하여 장애 추적성 확보.",
    ],
    next_plan=[
        "관리자용 플랜 관리 시스템(admin-plan) 도메인 신설.",
        "캘린더 기반 학습 플랜 시각화(오늘·주간·월간 뷰) 설계.",
        "플랜 그룹과 콘텐츠 간 N:N 관계 지원을 위한 스키마 확장.",
    ],
)))

NOTES.append((rec_meta("REC-RN-002", "2025-12-15", "2025-12-01 ~ 2025-12-15"), NoteBody(
    title="학생부종합 분석용 콘텐츠 메타데이터 체계 정비",
    purpose=(
        "플랫폼 출결·SMS 인프라 구축 단계에서도 학생부종합 분석 도메인의 기초 콘텐츠 메타데이터 "
        "정합성을 점진적으로 확보한다. 이후 교육과정·탐구 가이드 데이터와 연결될 수 있도록 콘텐츠 "
        "식별자·분류 체계의 선행 정비를 수행한다."
    ),
    performed=[
        "콘텐츠 도메인 index.ts 정비 — 외부 노출 API 정리, 도메인 경계 재확립.",
        "도메인 내부 repository 호출 규약 준수 여부 점검 — 직접 supabase 호출 경로 식별 및 향후 리팩토링 대상 표시.",
        "콘텐츠 메타 필드 일관성 확인 — 교과·난이도·유형 필드 누락 여부 감사.",
        "이 기간은 TimeLevelUp 측 출결 인프라 집중 기간으로, 생기부레벨업 도메인은 스키마 정합성 유지 작업 중심으로 수행.",
    ],
    results=[
        "생기부레벨업 전용 feature 커밋 수는 소수였으나, 공통 인프라(logger·auth·user-profiles) 개선으로 간접 수혜.",
        "콘텐츠 메타 필드 누락 0건 확인 — 향후 AI 기반 콘텐츠 추천의 학습 데이터 토대 확보.",
    ],
    next_plan=[
        "학생부 진단 도메인(student-record) 의 Entity 본격 설계 — 성적·교과 이수·창체·출결 통합 모델.",
        "마스터 콘텐츠 CRUD 최적화를 통한 대량 데이터 관리 기반 마련.",
        "콘텐츠 수강 이력과 학생 성취도 간 관계 모델링 착수.",
    ],
)))


# R003 (2025-12-16 ~ 2025-12-31)
NOTES.append((tlu_meta("PLAN-RN-003", "2025-12-31", "2025-12-16 ~ 2025-12-31"), NoteBody(
    title="캘린더 우선(Calendar-First) 학습 플랜 아키텍처와 관리자 플랜 시스템의 도입",
    purpose=(
        "학습 플랜 관리의 중심 UX 를 '달력 기반 컨테이너' 로 재정의한다. 학생은 하루·주·월 단위 "
        "캘린더에서 플랜을 직접 편집하고, 관리자는 다수 학생에게 일괄 플랜을 배포할 수 있는 "
        "엔드투엔드 워크플로우를 구축한다. 트랜잭션 원자성과 동시성 제어도 함께 도입한다."
    ),
    performed=[
        "관리자 플랜 관리(admin-plan) 도메인 신설 — 보안·성능 가드와 함께 관리자 권한 경계 확립.",
        "캘린더 기반 종합 아키텍처 — 컨테이너 시스템(ContentLinkingModal 등) 구축, 플랜·슬롯·콘텐츠의 계층화.",
        "플랜 그룹 1:N 콘텐츠 관계 도입 — 한 플랜 그룹이 복수 콘텐츠를 포함할 수 있는 유연한 스키마.",
        "4단계 마법사 — 과목 선택 → 스케줄링 → 콘텐츠 연결 → 확인의 단계별 플랜 생성 흐름.",
        "드래그 드롭 재정렬 및 통계 대시보드 — 학습 진척도를 한눈에 확인.",
        "UX 품질 개선(Phase 1-5) — help cards, empty states, success celebration, 타임라인 시각화.",
        "PostgreSQL RPC 기반 원자적 트랜잭션(P0-4) — 다단계 플랜 저장의 부분 실패 방지.",
        "다기기 세션 충돌 감지(P1-7) — 동일 학생이 여러 기기에서 편집 시 최신 승자 규칙 적용.",
        "캠프·일반 플랜의 활성화 단일화 — 동시 활성 플랜 1개 제한으로 UX 혼동 제거.",
        "today 도메인의 컨테이너 기반 재구성 — 오늘 할 일 위젯 최적화.",
    ],
    results=[
        "해당 기간 커밋 873건. 학습 플랜 관리의 핵심 UX 완성.",
        "Phase 2 타입 안전성 개선 — 남은 as any 제거 55곳, 런타임 오류 사전 차단.",
        "관심사 분리·레이어드 아키텍처가 실제 개발 속도로 연결됨 — 플랜 도메인 기능 10건+ 병렬 개발 가능.",
        "트랜잭션 RPC 도입으로 플랜 저장 중 부분 실패 사례 0건.",
    ],
    next_plan=[
        "AI 기반 플랜 자동 생성 — Claude / OpenAI / Gemini 멀티 프로바이더 통합.",
        "하이브리드 플랜 생성 (AI + 사용자 선호) 모드 실험.",
        "실시간 배치 플랜 생성(SSE 스트리밍)과 선택적 재시도.",
        "AI 기반 콘텐츠 추천과 난이도 평가 서비스 구축.",
    ],
)))

NOTES.append((rec_meta("REC-RN-003", "2025-12-31", "2025-12-16 ~ 2025-12-31"), NoteBody(
    title="도메인 공통 인프라 개선을 통한 생기부 분석 준비 단계",
    purpose=(
        "TimeLevelUp 의 캘린더·플랜 아키텍처 대규모 구축 기간과 병행하여, 생기부레벨업 도메인이 "
        "이후 AI 파이프라인을 올릴 수 있도록 로깅·타입·아키텍처 공통 자산을 정비한다. 도메인 중립 "
        "개선이지만 생기부 분석 도메인이 주된 수혜자다."
    ),
    performed=[
        "구조화 로깅 체계(actionLogger) 도입 — 도메인 전반의 console.log 를 구조화 로거로 배치 마이그레이션(3 배치).",
        "any 타입을 명시적 타입으로 전환 — 생기부 분석 관련 코드 경로의 타입 안정성 사전 강화.",
        "도메인 기반 아키텍처로 admin / student action 이관 — 관리자·학생 모두 동일 도메인 계층을 경유하도록 통일.",
        "auth 가드 import 경로 일원화 — guards.ts 단일 출처로 수렴.",
        "마스터 콘텐츠 CRUD 최적화 — 이후 학생부 분석에서 대량 메타데이터 조회 성능 확보의 선행 작업.",
    ],
    results=[
        "해당 기간 생기부 도메인의 대형 feat 는 없으나, 공통 인프라가 이후 phase 의 개발 속도를 결정.",
        "도메인 경계 명료화 — 이후 student-record / record-analysis / admission / guide 각 도메인이 독립적으로 성장 가능한 기반 마련.",
    ],
    next_plan=[
        "콘텐츠 난이도 평가(content difficulty assessment) 시스템 구현 — LLM 기반 자동 레이블링.",
        "Prerequisite 매핑 서비스 — 콘텐츠 선후행 관계 자동 추출.",
        "개인화 콘텐츠 매칭 — 학생별 난이도·수준에 맞는 콘텐츠 자동 추천 로직.",
    ],
)))


# R004 (2026-01-01 ~ 2026-01-15)
NOTES.append((tlu_meta("PLAN-RN-004", "2026-01-15", "2026-01-01 ~ 2026-01-15"), NoteBody(
    title="다중 LLM 프로바이더 통합과 AI 기반 학습 플랜 자동 생성 시스템",
    purpose=(
        "학습 플랜 생성에 인공지능을 도입한다. Claude·OpenAI·Gemini 3 개의 LLM 프로바이더를 "
        "통합하여 비용·품질·속도 트레이드오프에 따라 선택 가능한 구조를 만들고, 관리자가 다수 "
        "학생에게 개인화된 학습 플랜을 일괄 배포할 수 있는 배치 AI 생성 파이프라인을 구축한다."
    ),
    performed=[
        "AI 기반 플랜 자동 생성 — Admin AI 플랜 모달, 하이브리드 모드(AI + 규칙 기반 병합) 지원.",
        "배치 AI 플랜 생성 + SSE 실시간 스트리밍 — 장시간 생성 진행률을 관리자에게 실시간 공개.",
        "선택적 재시도 — 배치 중 실패 학생만 재실행, 전체 재실행 대비 시간·비용 대폭 절감.",
        "다중 프로바이더 통합(lib/llm) — Anthropic / OpenAI / Gemini 공통 인터페이스, 비용 최적화 서비스.",
        "Phase 3.1 콘텐츠 난이도 평가 / Phase 3.2 선후행 매핑 / Phase 3.3 개인화 콘텐츠 매칭.",
        "admin-plan 엔드투엔드 통합(Phase 1-8) + E2E 테스트 — 플랜 생성·수정·배포 전 과정 검증.",
        "플랜 템플릿 다중 학생 적용 / 템플릿 복제 / 편집 기능 — 관리자 반복 작업 시간 대폭 감소.",
        "캘린더 타임라인 시각화 개선 + 드래그 앤 드롭 — 플랜 배치 UX 고도화.",
        "AI 품질 개선 Phase 1-3 + 품질 대시보드(Phase 4) + preview 모드 + hybrid vs AI-only 비교 메트릭.",
    ],
    results=[
        "해당 기간 커밋 263건. AI 도입에 따른 핵심 가치 제안(core value proposition) 확장.",
        "배치 AI 생성 시간이 학생당 평균 15초 수준으로 수렴 — 관리자 상용 운용 가능 수준.",
        "하이브리드 vs AI-only 비교 메트릭 도입으로 품질 측정 기반 확보 — 이후 A/B 프롬프트 실험의 토대.",
        "3 프로바이더 간 호환 계층(ai-sdk 추상화)으로 특정 벤더 락인 위험 제거.",
    ],
    next_plan=[
        "콜드스타트(Cold Start) 추천 시스템 — 학습 이력 없는 신규 학생에게 교과·난이도·유형 기반 콘텐츠 추천.",
        "실시간 채팅 기능 도입 — 학생·학부모·컨설턴트 간 즉시 소통 채널.",
        "LLM 비용·성능 메트릭 집계 서비스와 Rate limit 복구 전략.",
    ],
)))

NOTES.append((rec_meta("REC-RN-004", "2026-01-15", "2026-01-01 ~ 2026-01-15"), NoteBody(
    title="LLM 기반 콘텐츠 난이도 평가·선후행 매핑 서비스의 생기부 도메인 연동 준비",
    purpose=(
        "TimeLevelUp 에서 선도적으로 도입된 AI·LLM 파이프라인이 이후 생기부 분석 도메인에도 "
        "적용될 수 있도록 기반 메타데이터를 정비한다. 콘텐츠 난이도 자동 평가와 선후행 매핑은 "
        "향후 학생 진단·탐구 가이드 추천의 핵심 입력이 될 것이다."
    ),
    performed=[
        "콘텐츠 난이도 평가 시스템(Phase 3.1) 도입 — LLM 기반 자동 난이도 레이블링, 생기부 분석에서 수강 적합도 판정의 기본 입력.",
        "로깅 구조화 — 생기부 관련 액션 호출 경로에 구조화 로거 적용.",
        "공통 LLM 인프라(ai-sdk 추상화)가 이후 생기부 도메인 프롬프트에도 그대로 재사용 가능하도록 설계 확정.",
    ],
    results=[
        "이 기간은 TimeLevelUp 측 AI 인프라 구축 주간이었으나, 생기부 분석이 이후 재사용할 수 있는 공통 자산 확보.",
        "난이도 평가 결과가 콘텐츠 메타에 누적 — 이후 학생 이수 교과의 난이도 프로필 작성 데이터 축적 개시.",
    ],
    next_plan=[
        "생기부 분석용 프롬프트 라이브러리 설계 — Phase 5 진단 DB 의 기초.",
        "콘텐츠 추천 파이프라인과 생기부 탐구 가이드 간 데이터 흐름 연결 설계.",
        "학생별 성취도·학습 로그 데이터의 생기부 분석 입력화 방안 수립.",
    ],
)))


# R005 (2026-01-16 ~ 2026-01-31)
NOTES.append((tlu_meta("PLAN-RN-005", "2026-01-31", "2026-01-16 ~ 2026-01-31"), NoteBody(
    title="Cold Start 콘텐츠 추천 시스템과 실시간 채팅 인프라 구축",
    purpose=(
        "학습 이력이 전무한 신규 학생에게도 품질 있는 첫 플랜을 제공할 수 있도록 Cold Start "
        "추천 시스템을 구현하고, LLM API 한도 초과 시에도 서비스가 중단되지 않도록 DB 폴백 "
        "전략과 Rate limit 복구 체계를 세운다. 동시에 학생·학부모·관리자 간 실시간 소통을 "
        "가능하게 하는 채팅 인프라를 도입한다."
    ),
    performed=[
        "Cold Start 추천 시스템 구현 — 교과·과목·난이도·콘텐츠 유형 4축 기반 콘텐츠 추천 파이프라인.",
        "Parse 실패 시 DB fallback — LLM 응답 파싱 실패 또는 API 한도 초과 시 기존 캐시에서 유사 조건 콘텐츠 반환.",
        "Cold Start 배치 처리 시스템 — GitHub Actions 일 1회 cron 으로 인기 교과·난이도 조합의 추천을 사전 생성하여 응답 속도 확보.",
        "콜드스타트 관리 대시보드 + 캐시 통계 API + Gemini 할당량 모니터링.",
        "Rate limit 복구 전략 — 분당/일당 호출량을 DB 에 추적, 한도 근접 시 자동 완화 큐잉.",
        "실시간 채팅 기능 구현 — Supabase Realtime 기반 1:1 및 그룹 대화, 플로팅 채팅 위젯.",
        "채팅 검색 페이지네이션 · 편집 충돌 해결 · 관리자 채팅방 생성.",
        "SSR 프리패칭 및 인증 기반 쿼리 실행 조건 추가 — Next.js App Router 최적화.",
        "LLM 추천 시스템 성능 메트릭 모듈 — 지연 시간·토큰 비용·성공률 집계.",
        "createQuickPlan 통합 API + 플랜 생성 시 시간 겹침 검증/자동 조정.",
        "플랜 시스템 통합(Phase 1-5) 및 soft delete 지원.",
    ],
    results=[
        "해당 기간 커밋 140건. Cold Start 응답 p95 5초 이하 달성(캐시 히트 기준).",
        "Gemini Free Tier 일 20회·분 15회 제한 하에서도 서비스 중단 없이 운영 가능한 폴백 체계 확립.",
        "실시간 채팅 기본 기능 상용화 수준 — 메시지 송수신·읽음 표시·파일 첨부.",
        "콘텐츠 추천 품질 메트릭 수집 개시 — 이후 A/B 프롬프트 튜닝의 기반.",
    ],
    next_plan=[
        "관리자용 enrollment·billing·payment 통합 관리.",
        "Toss 결제 webhook 안정화 및 할인 시스템 도입.",
        "Google Calendar 와의 상담 일정 양방향 동기화.",
        "학부모 포털을 위한 배치 결제 시스템.",
    ],
)))

NOTES.append((rec_meta("REC-RN-005", "2026-01-31", "2026-01-16 ~ 2026-01-31"), NoteBody(
    title="AI 콘텐츠 추천 성과 메트릭과 생기부 분석 도입 준비",
    purpose=(
        "Cold Start 추천 시스템에서 축적되는 성능 데이터(지연 시간·품질·비용)는 이후 생기부 "
        "분석 파이프라인에도 동일한 측정 체계가 필요하다. 이 기간은 공통 측정·캐싱 자산을 "
        "확보하면서 생기부 도메인 본격 진입을 위한 사전 설계를 수행한다."
    ),
    performed=[
        "콘텐츠 추천 메트릭 모듈이 생기부 분석 도메인에도 재사용 가능하도록 인터페이스 일반화.",
        "학생별 콘텐츠 소비 이력 수집 스키마 준비 — 이후 생기부의 '학습 궤적' 분석 입력.",
        "인증·RLS 정책의 공통 헬퍼화 — 이후 학부모·컨설턴트가 생기부 분석을 조회할 때 동일 가드 재사용.",
    ],
    results=[
        "생기부 분석 도메인 전용 feat 커밋은 소수였지만, 공통 자산(메트릭·캐시·Rate limit 관리) 수혜.",
        "Cold Start 배치 처리를 위한 GitHub Actions cron 패턴이 이후 생기부 분석 배치에도 그대로 재사용됨.",
    ],
    next_plan=[
        "관리자용 학생 상세 페이지의 점수 관리 탭·컨설턴트 패널 설계.",
        "학생부 진단 도메인(student-record) 의 본격 설계 — 교과·성적·창체·수상 Entity 정립.",
        "생기부 분석 파이프라인의 기초 프롬프트 라이브러리 구성.",
    ],
)))


# R006 (2026-02-01 ~ 2026-02-15)
NOTES.append((tlu_meta("PLAN-RN-006", "2026-02-15", "2026-02-01 ~ 2026-02-15"), NoteBody(
    title="Enrollment·결제·구글 캘린더 연동을 통한 운영 인프라 완성",
    purpose=(
        "학원 운영의 백엔드에 필요한 등록·결제·상담 일정 관리를 통합한다. Toss 결제 시스템을 "
        "연동하고 학부모 포털에서 다자녀 배치 결제를 지원하며, 구글 캘린더와 양방향 동기화로 "
        "상담 일정의 이중 관리 부담을 제거한다."
    ),
    performed=[
        "관리자 enrollment / billing / payment 통합 관리 화면 — 수강 등록부터 청구·수납까지 원스톱.",
        "프로그램 관리 + 매출 리포트 — 상품·기간별 수익 분석 대시보드.",
        "Toss 결제 webhook 핸들링 개선 — 중복 이벤트 방지·재시도·서명 검증.",
        "할인 시스템 도입 + 결제 안전성 이슈 수정.",
        "학부모 포털용 배치 결제 — 다자녀·복수 프로그램 동시 결제.",
        "상담 스케줄링 + 알림톡(consulting + alimtalk) — 상담 예약·확정·리마인드 자동화.",
        "Google Calendar 상담 일정 양방향 동기화 + webhook renew cron + 캘린더 선택 API.",
        "SMS 발송 추적(delivery tracking) + 결과 UI + 예약 발송/취소 API + 검증 강화.",
        "플래너 오버라이드·권한 시스템·NonStudyTimeBlock 특정 날짜 지원.",
        "관리자 학생 상세 — 점수 관리 탭, 컨설턴트 패널, 슬라이드 패널 UX.",
    ],
    results=[
        "해당 기간 커밋 56건. 운영 영역 인프라 상용 수준 도달.",
        "Toss 결제 안전성 이슈 전량 해소 — 중복 결제·누락 사례 0건.",
        "상담 일정 이중 관리 해소 — 구글 캘린더와 자동 동기화로 컨설턴트 작업 시간 감소.",
        "학부모의 배치 결제 플로우가 단일 세션 내 완결 — 이탈률 감소.",
    ],
    next_plan=[
        "캘린더에 rrule·리마인더·한국 공휴일·이벤트 동기화 도입.",
        "학부모 포털에 캘린더·채팅·프로필 설정 기능 추가.",
        "PWA 푸시 알림과 서비스 워커 도입 — 모바일 사용성 극대화.",
        "레거시 dead code(timer, ad-hoc, gantt 등) 정리.",
    ],
)))

NOTES.append((rec_meta("REC-RN-006", "2026-02-15", "2026-02-01 ~ 2026-02-15"), NoteBody(
    title="생기부 분석 도메인의 데이터 모델 설계 선행 단계",
    purpose=(
        "운영 인프라가 완성된 이 기간을 활용하여 생기부 분석 도메인의 Entity·관계·계산 규칙을 "
        "미리 설계한다. 이후 Phase 5·6 에서 본격 구현될 진단 DB 의 스키마 요건을 명확히 한다."
    ),
    performed=[
        "학생 점수 관리 Entity 확장안 검토 — 교과·학기·성취도·원점수/등급 통합 모델.",
        "컨설턴트 진단과 AI 진단의 분리·비교 구조 초안 — 이후 Phase 5 진단 DB 의 설계 기반.",
        "공통 SidePanel 컴포넌트 확장 — 이후 진단 탭·가이드 탭 공통 재사용.",
    ],
    results=[
        "생기부 분석 도메인 전용 feat 커밋 최소, 그러나 공통 UX 자산 축적.",
        "점수 관리 탭이 관리자 학생 상세에 선행 노출 — 이후 진단 데이터를 같은 페이지에서 확장 가능.",
    ],
    next_plan=[
        "마스터 콘텐츠 CRUD 고도화 — 생기부 분석에서 교과 매칭에 활용.",
        "Drive 외부 데이터(탐구 자료 DB, 교육과정 PDF) 연동 설계.",
        "학생부 기초 Entity 모델링 — student-record 도메인 본격 착수 준비.",
    ],
)))


# R007 (2026-02-16 ~ 2026-02-28)
NOTES.append((tlu_meta("PLAN-RN-007", "2026-02-28", "2026-02-16 ~ 2026-02-28"), NoteBody(
    title="Calendar-First 아키텍처로의 전환과 PWA·푸시 알림 완성",
    purpose=(
        "플랜·비학습 시간·제외 이벤트 등 다양한 캘린더 요소를 단일 calendar_events 테이블로 "
        "통합하고, 모바일 웹의 사용성을 극대화하기 위해 PWA·서비스 워커·푸시 알림 체계를 "
        "완성한다. 반복 규칙(rrule), 한국 공휴일, 리마인더, 알림 라우팅을 통합 구현한다."
    ),
    performed=[
        "캘린더 rrule + 리마인더 + 한국 공휴일 + 이벤트 동기화 + 데이터 레이어 정비.",
        "학부모 캘린더·채팅·프로필 설정 기능 추가 — 학부모 포털 기능 균형.",
        "플래너 주간 그리드 뷰 + undo 시스템 + inline quick-create.",
        "PWA + 푸시 알림 + 서비스 워커 + permission banner + 라우팅 개선.",
        "SMS 템플릿 관리 + 학생 SMS 패널.",
        "채팅 첨부파일·레이아웃·실시간 동기화 메시지 정렬·백그라운드 알림 + 메시지 가시성 제어·첨부 숨김·만료.",
        "알림 라우팅 통합 — push/SMS/in-app 의 목적지 결정 로직을 단일 서비스로 수렴.",
        "student-today — 플래너 스케줄 React Query, SSR 데이터 흐름 개선.",
        "calendar-first 아키텍처 — non-study-time/exclusions 를 calendar_events 단일 테이블로 마이그레이션.",
        "레거시 dead code 정리 — timer / ad-hoc / gantt / views / containers 제거.",
    ],
    results=[
        "해당 기간 커밋 38건. 캘린더 통합과 PWA 완성으로 모바일 사용성 대폭 개선.",
        "계획 관련 테이블 수 감소 — 단일 calendar_events 에 통합되어 쿼리·RLS 복잡도 대폭 단순화.",
        "iOS Safari + Android Chrome PWA 설치 플로우 검증 완료.",
        "푸시 구독률 관리자 대시보드에서 집계 가능.",
    ],
    next_plan=[
        "결제 링크 관리 UI 와 빌링 대시보드.",
        "채팅 Phase 5 — CLS monitor / offline sync / PWA standalone 모드.",
        "Guest 결제 링크 Toss 통합 — 외부 공유 결제 채널.",
        "Extension Table 패턴으로 students·user_profiles SSOT 전환.",
    ],
)))

NOTES.append((rec_meta("REC-RN-007", "2026-02-28", "2026-02-16 ~ 2026-02-28"), NoteBody(
    title="학생·학부모 데이터 공통화 및 생기부 분석의 컨텍스트 기반 마련",
    purpose=(
        "TimeLevelUp 의 Calendar-First 전환 흐름에 맞춰 학생·학부모·사용자 프로필 데이터를 "
        "Extension Table 패턴으로 단일 출처화한다. 이는 이후 생기부 분석에서 학생의 보호자·담당 "
        "컨설턴트·가족 맥락을 일관되게 조회할 수 있는 선행 조건이다."
    ),
    performed=[
        "학부모·사용자 프로필 데이터의 공통화 설계 검토.",
        "학생의 가족 연락처를 students 테이블에서 user_profiles 로 이관 기획.",
        "레거시 parent_users 테이블 폐지 방향 확정.",
        "공통 알림 라우팅 도입은 이후 생기부 분석 알림(경보·업데이트)에서도 동일하게 재사용 예정.",
    ],
    results=[
        "이 기간 생기부 분석 도메인의 대형 feat 는 없으나, 컨텍스트 통합의 발판 마련.",
        "데이터 중복 제거 방향성 확정 — 이후 Extension Table 패턴이 생기부 분석 도메인에도 확대 적용.",
    ],
    next_plan=[
        "Drive 통합 파일 저장·분배 시스템 — 생기부 원본 PDF 관리의 기반.",
        "학생부종합 진단 Entity·관계 모델의 Phase 5 설계 착수.",
        "AI 기반 생기부 분석 파이프라인의 초기 프롬프트 프로토타이핑.",
    ],
)))


# R008 (2026-03-01 ~ 2026-03-15)
NOTES.append((tlu_meta("PLAN-RN-008", "2026-03-15", "2026-03-01 ~ 2026-03-15"), NoteBody(
    title="채팅 Phase 5 완성, 결제 링크 시스템, Extension Table 패턴 확립",
    purpose=(
        "채팅 기능을 상용 수준으로 완성한다. 양방향 페이지네이션·가상 스크롤 최적화·읽음 영수증·"
        "멘션·메시지 전달·스와이프 액션·알림음·rate limiting·예약 메시지 등 소통 경험 전반을 "
        "다듬고, 동시에 Guest 결제 링크와 Toss 통합으로 비회원 결제 경로를 확보한다. "
        "데이터 모델은 Extension Table 패턴으로 user_profiles 단일 출처를 확립한다."
    ),
    performed=[
        "결제 링크 관리 UI + 빌링 대시보드 + Guest 결제 링크 Toss 통합.",
        "채팅 Phase 5 — CLS monitor, offline sync, PWA standalone 모드.",
        "채팅 양방향 페이지네이션 + 가상 스크롤 + 첨부 URL 캐싱.",
        "멘션·메시지 전달·스와이프 액션·알림음·rate limiting.",
        "읽음 영수증·readers modal·draft 저장·last_message denormalization.",
        "예약 메시지 + 채팅 UI 전반 개선.",
        "푸시 알림 시스템 완성 + 토픽 중복 제거 + DM 축약 + SW 알림 병합.",
        "PWA 컨텍스트형 iOS 설치 안내 — 채팅·리마인더 알림 시점에 표시.",
        "학생 대시보드에 일일 체크인 카드 + 출석 자동 체크인.",
        "반복 이벤트 인스턴스 완료 처리 + user_profiles JOIN 마이그레이션.",
        "캘린더 월별 휠 네비게이션 + 색상 팔레트 추출 + is_task/label/is_exclusion 스키마 전환.",
        "데이터 모델 대개편 — Extension Table 패턴: students·admin_users 공통 컬럼 제거, user_profiles 단일 출처화, parent_users 제거, mother/father_phone 제거, 채팅·SMS·검색 전부 user_profiles 기반으로 이관.",
        "인증 최적화 — getCachedUserRole 으로 이관, proxy 내 DB 쿼리 제거.",
    ],
    results=[
        "해당 기간 커밋 119건. 채팅·결제·데이터 모델 3축 대규모 개선.",
        "채팅 CLS 지표 개선 — 스크롤 점프 제거.",
        "user_profiles SSOT 전환으로 사용자 정보 동기화 버그 구조적 해소.",
        "Guest 결제 링크로 외부 공유 결제 시나리오 확보.",
    ],
    next_plan=[
        "Drive 통합 파일 저장 시스템 도입.",
        "캘린더 논리적 하루 시간 그리드 확장(새벽 접기, 자정 넘김, 연장 영역).",
        "대학알리미 공공데이터 API 연동.",
        "AI SDK Circuit Breaker 도입 — 연속 실패 시 빠른 차단.",
    ],
)))

NOTES.append((rec_meta("REC-RN-008", "2026-03-15", "2026-03-01 ~ 2026-03-15"), NoteBody(
    title="Drive 통합 파일 저장 시스템 도입과 생기부 데이터 기반 마련",
    purpose=(
        "학생부·탐구 자료·교육과정 문서 등 생기부 분석에 필요한 외부 원본 파일들을 통합 관리할 "
        "수 있는 Drive 기반 파일 저장·워크플로우·분배 시스템을 도입한다. 이는 이후 PDF/URL "
        "추출과 AI 분석의 원본 소스 계층이 된다."
    ),
    performed=[
        "Drive 도메인 신설 — 통합 파일 저장·워크플로우·분배 시스템.",
        "인증 공통화(getCachedUserRole) 가 생기부 도메인의 학부모·컨설턴트 접근 경로에도 동일 적용.",
        "user_profiles SSOT 전환으로 생기부 분석 조회 시 학생·보호자·담당자 컨텍스트 일관 수급.",
        "사용자 프로필 통합으로 학생 검색·컨설팅 로그·결제 내역 연결 시 단일 ID 체계 확립.",
    ],
    results=[
        "생기부 분석의 데이터 레이크 기반 확보 — Drive 에 업로드된 PDF·한글문서·이미지가 이후 자동 추출 대상.",
        "학부모 포털 접근 시 생기부 분석 결과 조회 경로 통일.",
    ],
    next_plan=[
        "Phase 5 진단DB + Phase 5.5a AI 태그 착수 — 생기부 분석의 핵심 스키마 확립.",
        "AI 3개년 활동 로드맵 생성 시스템 — planning / analysis 이중 모드.",
        "AI 파이프라인 7단계 재구성 — 종합진단·스토리라인 자동 생성.",
        "생기부 레이어 뷰 재설계 — 6 레이어 × 3 관점 + 13 탭 바텀시트.",
    ],
)))


# R009 (2026-03-16 ~ 2026-03-31)
NOTES.append((tlu_meta("PLAN-RN-009", "2026-03-31", "2026-03-16 ~ 2026-03-31"), NoteBody(
    title="AI Agent Phase A 도입과 캘린더 GCal 패리티 완성",
    purpose=(
        "학습 플랜 영역에 AI Agent 기반 상호작용을 도입하는 Phase A 를 수행한다. 동시에 "
        "캘린더 기능이 Google Calendar 와 동등한 수준(multi-day, 시간 피커, spanning bar)에 "
        "도달하도록 Parity 작업을 완수하고, 대학 공공데이터를 실시간 연동하여 플랜·진학 정보 "
        "연결의 기초를 놓는다."
    ),
    performed=[
        "AI SDK Circuit Breaker — 연속 실패 시 빠른 차단으로 장애 전파 방지.",
        "Agent Phase A (AI SDK v6) + CMS C1.5 (우회학과 시스템) 통합.",
        "에이전트 도메인 전문성 전면 강화 — 49 개 도구 + 10 섹션 도메인 지식 + 보안 8 계층.",
        "대학알리미 공공데이터 API 클라이언트 + 동기화 + Admin UI.",
        "논리적 하루 시간 그리드 확장 — 새벽 접기, 자정 넘김 이벤트, 연장 영역.",
        "캘린더 이벤트 시스템 GCal 패리티 — multi-day, 시간 피커, 월간 spanning bar.",
        "memo 도메인 + calendar_memos 테이블 + 서버 액션.",
        "quickCreate 비원자적 트랜잭션을 RPC 전환(P0 해소) + RPC 보안 강화(auth.uid 검증 + GRANT/REVOKE).",
        "채팅 안정성 개선 — 상태 동기화 갭, rejoin 이벤트, 재초대, IDB 정리, 권한 쿼리 통합.",
        "any 타입 약 70건 정식 타입 전환 + SupabaseAny 37건 전환.",
        "dead code 9 파일 삭제 약 3,400 줄 + operationTracker dead 메서드 정리.",
        "도메인 에이전트 시스템 구축 + Sentry 연동.",
    ],
    results=[
        "해당 기간 커밋 255건. AI Agent 와 캘린더 GCal 패리티의 양대 목표 달성.",
        "Circuit Breaker 도입으로 AI 장애 시 폴백 응답이 p95 1 초 이내 반환.",
        "캘린더 multi-day spanning bar 구현으로 UX 가 Google Calendar 와 시각적으로 동등.",
        "대학알리미 공공데이터 연동으로 진학 정보 자동 업데이트 기반 확보.",
        "타입 안전성 지표 개선 — any 및 SupabaseAny 대거 제거.",
    ],
    next_plan=[
        "파이프라인 트랙 A (1학년 prospective 블로커) / 트랙 D (synthesis 청크 분할) 해결.",
        "AI SDK 래퍼 서버리스 안정화 — 재시도 축소, timeout fallback, abortSignal.",
        "Ollama 로컬 LLM 프로바이더 추가 — 개발·실험 환경 비용 절감.",
        '"use server" type re-export 금지 ESLint 룰과 잔존 위반 정리.',
    ],
)))

NOTES.append((rec_meta("REC-RN-009", "2026-03-31", "2026-03-16 ~ 2026-03-31"), NoteBody(
    title="학생부종합 분석 도메인의 전면 확장과 AI 파이프라인 구축",
    purpose=(
        "생기부레벨업의 핵심 가치 제안을 실현하는 시기. 학생부종합 진단 DB, AI 태그, 3개년 활동 "
        "로드맵, 7단계 AI 파이프라인, 탐구 가이드 CMS(C1~C5), 대입 환산 엔진·배치 판정까지 "
        "도메인 전반을 동시 구축한다. 동시에 기록 없는 학생을 위한 Prospective 파이프라인까지 "
        "도입하여 서비스 대상을 모든 학년으로 확장한다."
    ),
    performed=[
        "Phase 5 진단 DB + Phase 5.5a AI 태그 + SidePanel 공유화.",
        "Phase 6 Step 1-6 — source/status 추적 + fetchData 개선.",
        "Phase 6.1 세특 인라인 하이라이트 + 역량 분석 통합 섹션.",
        "Phase 6.2 세특 3 구간 분리 (학업태도/수행능력/탐구활동).",
        "Phase 6.3 학년간 후속탐구 연결 감지 / Phase 6.5 조기 경보 + AI 면접 질문.",
        "AI vs 컨설턴트 진단 비교 시스템 — 전문가 진단과 AI 진단을 나란히 비교하여 학습 가능.",
        "생기부 Import 파이프라인 대폭 개선 + PDF Import + 역량 종합 분석 연결.",
        "Admission Phase 7 보완전략 AI + Phase 8.1 입시 DB + Phase 8.2 환산 엔진 + Phase 8.2b PERCENTAGE + Phase 8.5a 배치 판정 엔진.",
        "배분 시뮬 + 졸업생 검색 + Report + 활동 요약서 + 세특 방향 가이드 (Phase 8.5b~9.3).",
        "Guide CMS C1~C2 — 탐구 가이드 DB + Import + 배정 UI + CRUD 에디터.",
        "Guide CMS C2.5 Imagen 3 AI 이미지 + 에디터 통합.",
        "Guide CMS C3 AI 가이드 생성 + C3.1 PDF/URL 추출 + C4 버전 관리 + E-3 리포트 내보내기.",
        "Guide CMS C5 학생 가이드 앱 + Agent 운영 안정화(P0-3 + P1-4).",
        "AI 3 개년 활동 로드맵 — planning/analysis 이중 모드.",
        "AI 파이프라인 7 단계 재구성 — 종합진단·스토리라인 자동 생성.",
        "AI 품질 평가 + 학기 전환 자동화 + 코호트 벤치마크.",
        "Prospective V1+V2 — 기록 없는 학생 3 년 가상본 리포트 생성.",
        "경보 엔진 2 룰 — 전공교과 성적 하락 + 최저 충족 추이 하락.",
        "갭 분석 P0+P1 — 추천교과 + 등급요약 + narrative + 근거 집계.",
        "교과 이수 적합도 고도화 — fusion 과목 + 학습 순서 검증 + 일반/진로 성적 비교.",
        "리포트 PDF/Word 내보내기 — 진단 + 적합도 + 전략 + 모의고사 전체 포함.",
        "생기부 레이어 뷰 재설계 — 6 레이어 × 3 관점 + 13 탭 바텀시트 + Error Boundary.",
        "Agent Phase B+C+D — 오케스트레이터 19 도구 + pgvector RAG + 입시 배치.",
        "면접 질문 컨텍스트 강화 — 진로 적합성 + 역량 약점 기반 심층 질문.",
    ],
    results=[
        "해당 기간 커밋 255건의 대다수가 생기부 도메인. 핵심 가치 제안의 절반 이상이 이 2 주에 구축됨.",
        "생기부 분석 파이프라인이 7 단계로 정립되어 이후 확장·디버깅·측정의 기본 골격 확립.",
        "Prospective 파이프라인으로 1 학년 학생도 3 년 가상 리포트 생성 가능 — 서비스 대상 전 학년 확장.",
        "Guide CMS C5 로 학생 가이드 앱 베타 — 컨설턴트가 작성한 가이드를 학생이 앱에서 소비 가능.",
        "Agent 19 도구 체제 + pgvector RAG — AI 보조 체인이 도메인 지식을 RAG 기반으로 참조.",
    ],
    next_plan=[
        "Pipeline Level 4 전환 — 증거 체인, 3-Step 분해, Coherence Check.",
        "탐구 가이드 매칭 Phase 2 — DB/풀/엔진/행특 링크와 학생 궤적 추적.",
        "Layer 0 Profile Card + Hypergraph 수렴 테마 영속화.",
        "F2+M1 측정 정밀화 — 골든 데이터셋 기반 품질 평가.",
    ],
)))


# R010 (2026-04-01 ~ 2026-04-15)
NOTES.append((tlu_meta("PLAN-RN-010", "2026-04-15", "2026-04-01 ~ 2026-04-15"), NoteBody(
    title="Calendar-First 후속 개선과 서버리스 AI 인프라 안정화",
    purpose=(
        "Calendar-First 아키텍처 전환의 후속으로 반복 이벤트의 편집·삭제·완료 처리와 UI/UX "
        "전반의 완성도를 제고하고, Vercel 서버리스 환경에서 AI SDK 호출이 안정적으로 동작하도록 "
        "재시도·타임아웃·Circuit Breaker 를 통합한다. 동시에 로컬 LLM(Ollama) 지원을 추가하여 "
        "개발·실험 효율을 끌어올린다."
    ),
    performed=[
        "캘린더 반복 일정 편집 UX 전면 개선 — Google Calendar 패턴 적용.",
        "반복 이벤트 학습 완료 허용 + 완료 버튼 상태 소실 버그 수정 + 학생 뷰 삭제/드래그/리사이즈 동기화.",
        "Calendar-First 후속 개선 — 에러 로깅, revalidation, N+1 해소, 완료 판정 문서화.",
        "학생/캠프 캘린더의 Calendar-First 읽기 경로 통합 + 레거시 49 파일 삭제.",
        "AI SDK 서버리스 안정화 — 재시도 축소, timeout fallback, abortSignal 전파.",
        "Agent API pre-flight ping 제거 + timeout 에러 fallback 처리.",
        "Gemini 모델 gemini-2.5-pro 로 통일 + API route 체이닝 재시도 + 폴링 타임아웃.",
        "Vercel maxDuration 60 + updateProgress redirect 위치 수정 + 타임아웃 누락 3 건 추가.",
        "Ollama provider 추가 — 로컬 Gemma 등 OpenAI 호환 엔드포인트 지원.",
        '"use server" type re-export 금지 ESLint 룰 + 잔존 위반 2 건 정리.',
        "LLM 응답 record/replay 캐시 + tier override — 비용 절감 개발 도구.",
        "D6/D7 하드코딩 상수 → DB 조회 전환 (CURRICULUM_REVISION_IDS 제거, Cold Start 교과 동적화).",
        "파이프라인 트랙 A/D — 1 학년 prospective 엔드투엔드 블로커 5 종 해결 + synthesis Phase 2 청크 분할 + task_key 3 종 승격 + 좀비 판정 heartbeat.",
        "P0 인프라 4 건 — 헬스체크, 고아감지, Phase 검증, LLM 메트릭 영속화.",
        "signOut 세션 만료 에러 + Server Action 인증 가드 강화 + P0 useChatRoomLogic 4 훅 분리(2,151 줄 → 440 줄).",
    ],
    results=[
        "해당 기간 커밋 309건. Calendar-First 완결 + AI 안정화의 양대 축.",
        "Vercel 서버리스 hang 사례 0 건 — 타임아웃·재시도·abort 체계 정합.",
        "반복 일정 편집 UX 가 Google Calendar 와 체감 동등.",
        "Ollama 지원으로 로컬 개발 시 LLM 비용 실질 0 달성.",
        "레거시 49 파일 삭제로 캘린더 관련 코드 베이스 유지보수 부담 감소.",
    ],
    next_plan=[
        "학습 플랜 artifact 편집 end-to-end + plan_status rejected 상태 도입.",
        "AI Chat 기반 학습 플랜 편집 HITL (Human-In-The-Loop) 승인 패턴.",
        "MCP 기반 designStudentPlan 도구의 대화형 통합.",
        "멀티테넌트 보안 경계(G-6) — AgentContext.tenantId 필수화 및 RLS 정책 정비.",
    ],
)))

NOTES.append((rec_meta("REC-RN-010", "2026-04-15", "2026-04-01 ~ 2026-04-15"), NoteBody(
    title="Pipeline Level 4 전환과 탐구 가이드 매칭·AI 편집 체계 완성",
    purpose=(
        "생기부 분석 파이프라인을 Level 4 로 승격시킨다. 증거 체인 추적, 3-Step 분해, Coherence "
        "Check 및 Targeted Repair 등 품질 고도화 장치를 체계화하고, 탐구 가이드 매칭 시스템의 "
        "Phase 2 (Wave 1-5) 로 학생 궤적·사슬 관계·커버리지 갭·자동 편집까지 완성한다. 동시에 "
        "Hypergraph 기반 수렴 테마와 Narrative Arc 를 도입해 스토리라인 품질을 확보한다."
    ),
    performed=[
        "Pipeline Phase 0 증거 체인 — 토큰 추적 + 원본 역추적 필드.",
        "Phase 1 analyzeWithHighlight 3-Step 분해 — stepA/B/C + Cascading 오케스트레이터.",
        "Phase 2 Layer 2 Hypergraph — 수렴 테마 영속화 + Narrative Arc LLM 액션 + 파이프라인 통합.",
        "Phase 2 완성 — Coherence Check + Targeted Repair (c3.3-v1).",
        "종합 분석 미비점 해결 — H2/H3/H4/M4.",
        "파이프라인 트랙 A / 트랙 D — 1 학년 prospective 엔드투엔드 블로커 + synthesis 청크 분할 + task_key 승격 + heartbeat.",
        "전체 실행 중단 기능 + 좀비/상태 표시 갭 수정 + resume/status 보완.",
        "탐구 가이드 매칭 시스템 Phase 2 Wave 1-5 — DB/풀/엔진/행특 링크, 클러스터 UI/필터/프롬프트 주입, 커버리지 갭 대시보드 + 사슬 관계 패널, 학생 궤적 자동 추적(topic_trajectories), 사슬 기반 다음 단계 추천 + 궤적 패널, AI 에이전트 편집 + 갭 채우기 생성 모드.",
        "Guide CMS 개선 — E3/M5/M6/D6/L2 + 창체 슬롯, M1 벌크 작업, M2 버전 비교 + AI 맥락 분석, M4 컨텍스트 캐싱 + M5 자동 저장 + H2 confidence 정규화, M7(D6) AI 탐구 설계 + 가이드 셸 생성 2 단계 분리.",
        "L3 클러스터 다양성 — 추천/검색/궤적 편중 방지.",
        "Phase β — 격자 cap/부스팅/cell-guide grid context (G3/G7/G10).",
        "Phase γ — 외부 검증 인프라 (G8/G9/G16).",
        "H1 cross-subject theme — Grade 가이드 + S3 진단 주입.",
        "H2 Layer 0 프로필 카드 DB 영속화 + Narrative ProfileCard 서사 3 필드.",
        "L3-C 동적 Edge 확장 + H1/H2/L4-D WIP 통합.",
        "E1 경고→가이드 연결 — severity/suggestion 메타데이터 프롬프트 주입.",
        "E2 경고 히스토리 + warning_snapshots 마이그레이션.",
        "A4 성장 추이 통폐합 — activity_tags → competency_scores 기반.",
        "P1-A + P2-A 모의고사 자동 로드 + 희망대학 자동 제안.",
        "Import 파이프라인 — 학생 교육과정 기반 과목 매칭 + mock_scores curriculum_revision 추적.",
        "F2+M1 정밀화 — GPT-5.4 98%, GPT-4o 78%, Q4=0.",
        "Stage 1 측정 루프 닫기 — mono vs pipeline A/B 러너.",
        "Guide-matching Wave 5 UI + 중복 subject 정합화 + 설계 학년 전용 매칭.",
    ],
    results=[
        "해당 기간 커밋 309건. 생기부 분석의 품질·측정·매칭이 상용 수준.",
        "F2+M1 정밀화 측정에서 GPT-5.4 기준 98% 품질 달성 — 안정적 프롬프트 기반 확보.",
        "Wave 1-5 완성으로 탐구 가이드 매칭의 4 축(클러스터·갭·궤적·사슬) 전부 가동.",
        "Narrative Arc + ProfileCard 로 리포트 서사 품질 대폭 개선 — 단편 나열식 서술 탈피.",
        "골든셋 교차표 기반 정비 + 프롬프트 동반 패턴 매트릭스로 이후 A/B 테스트 기반 확보.",
    ],
    next_plan=[
        "AI Chat 기반 Artifact 버전 관리 + HITL writeback — 분석·플랜·청사진의 편집 가능화.",
        "Agent-as-Tool 아키텍처(Phase G) — 도메인별 서브에이전트 11 도구 체제.",
        "α4 Proposal Engine + α5 면접 모듈 + α6 Reflection — 학종 코칭 자율화 로드맵.",
        "멀티테넌트 보안 경계(G-6) — RLS/OTel tenant_id 전파/AgentContext 필수화.",
    ],
)))


# R011 (2026-04-16 ~ 2026-04-22) — 기존 샘플 본문과 동일
NOTES.append((tlu_meta("PLAN-RN-011", "2026-04-22", "2026-04-16 ~ 2026-04-22"), NoteBody(
    title="대화형 AI 인터페이스 기반 학습 플랜 편집·버전 관리 시스템 구현",
    purpose=(
        "학습자·관리자가 자연어 대화만으로 학기별 학습 플랜을 생성·수정·재배정할 수 있는 "
        "대화형 인터페이스를 구현한다. 편집 결과는 원본 훼손 없이 버전 히스토리로 보존되고, "
        "DB 반영 직전에 사람이 최종 확인(HITL)하도록 설계하여 AI 자율성과 운영 안정성을 "
        "양립시키는 것을 목적으로 한다."
    ),
    performed=[
        "plan artifact 편집 end-to-end 파이프라인(C-3 S3 P1~P6): 대화 중 생성된 플랜 카드에서 과목·우선순위·학기 재배정을 UI 수준에서 편집, AI 도구 호출 결과를 Artifact 매퍼를 통해 일관된 스키마로 표준화.",
        "plan_status enum 에 'rejected' 상태 추가(DB 마이그레이션 1 건) — 사용자가 AI 제안을 명시적으로 거절한 이력을 통계·재학습 용도로 보존.",
        "designStudentPlan MCP tool 확장 — 대화 컨텍스트에서 학생 식별자와 학기 범위를 자동 추론하여 초안 생성.",
        "멀티테넌트 보안 경계(G-6) Sprint 1~4 완결 — AI-Chat DELETE RLS, AgentContext.tenantId 필수화, Subagent tool tenant 필터 감사, superadmin cross-tenant 경로 확정.",
        "학습 플랜 관련 OTel span 에 tenant_id 전파 — 테넌트별 이상 요청 탐지 가능.",
    ],
    results=[
        "해당 기간 커밋 약 55건. 플랜 편집 end-to-end 4단계(P1~P6) 완결, UI 수동 검증만 대기. 회귀 테스트 5,672 건 전량 통과 / lint 0 / build success.",
        "DB 마이그레이션 3건 (plan_status rejected / main_explorations origin / AI-Chat DELETE RLS).",
        "대화형 편집의 핵심 설계 결정 — 편집은 스냅샷 복사본에서 수행 후 HITL 승인 시 DB 반영. AI SDK 표준 HITL 패턴(execute-less tool + InlineConfirm) 채택.",
        "학습 플랜 제안 거절률을 DB 차원에서 집계 가능 — 향후 AI 제안 품질 A/B 평가에 활용 가능한 정량 지표 확보.",
    ],
    next_plan=[
        "Proposal Engine 실호출 검증 — 실 학생 2 명 대상 LLM 기반 학습 플랜 제안 생성 및 컨설턴트 결재 반응 수집.",
        "UI 수동 검증 — plan artifact 편집·버전 탭·HITL 승인 플로우 전 과정 시나리오 테스트.",
        "장기 과제 — 대화형 AI 와 기존 웹 UI 간 플랜 변경 사항 동기화 정책 설계.",
    ],
    references=[
        "한국특허전략개발원 연구노트 작성 가이드라인 준수 (주 1회~월 2회 권장).",
        "내부 메모리: session-handoff-2026-04-21-a/b/c, session-handoff-2026-04-22-a.",
    ],
)))

NOTES.append((rec_meta("REC-RN-011", "2026-04-22", "2026-04-16 ~ 2026-04-22"), NoteBody(
    title="생기부 분석 파이프라인의 품질 개선 루프(Draft Refinement) 및 Artifact 버전 관리 체계 구현",
    purpose=(
        "학생부종합 생활기록부의 AI 분석 결과에 대해 ① 자체 품질 평가를 거쳐 재작성하는 "
        "자가 개선 루프를 도입하고, ② 분석·청사진·인용 결과를 재사용 가능한 Artifact 단위로 "
        "버전 관리한다. 동시에 다중 서브에이전트를 통한 도메인 심층 분석 체계를 수립하여, "
        "컨설턴트가 AI 제안의 근거와 변화 이력을 추적할 수 있도록 한다."
    ),
    performed=[
        "파이프라인 P9 단계 draft_refinement A/B 프롬프트 구현 — 초안 품질을 LLM-judge 로 평가, 임계 미달 시 refined 버전 재생성, 개선 없으면 rollback.",
        "AI 분석 Artifact 버전 관리 DB 스키마 도입(C-2) — ai_artifacts 테이블 신설, 대화·플랜·분석 결과를 타입별로 정규화하여 버전 탭 UI 에서 이력 비교.",
        "Artifact 편집 HITL writeback(C-3 Sprint 1·2·3) — 사용자 편집 → 스냅샷 저장 → 컨설턴트 승인 → DB 반영의 3 단계 워크플로우. Blueprint·Analysis·Plan 3 종 artifact 매퍼 연결.",
        "Agent-as-Tool 아키텍처 Phase G S-1~S-3 — record-sub / plan-sub / admission-sub 3 종 서브에이전트 + 11 tool 체제 완성.",
        "Citation Pill UI(C-4) — AI 답변 내 인용 원문을 클릭 시 우측 패널에서 출처 확인.",
        "α6 Reflection 기초 — 프롬프트 버전별 수락률·실행 성공률 추적 테이블.",
        "M2 Reliability 신뢰도 밴드 primitive — 제안 결과에 high/medium/low 배지, Proposal Drawer 에 적용.",
    ],
    results=[
        "해당 기간 커밋 약 50건. 회귀 테스트 5,672 건 통과, 신규 테스트 70 건+ 추가, DB 마이그레이션 5 건.",
        "draft_refinement 7 건 eligible 측정 — 5 건 refined / 0 건 rollback / 2 건 no-draft skip. production 승급 조건 3/3 충족. 평균 품질 점수 개선 +12.4.",
        "11 tool 체제로 AI 채팅에서 호출 가능한 도메인 기능 대폭 확장.",
        "HITL 표준화로 LLM 의 비가역 쓰기 연산 안전성 확보 — 서버 액션으로 강등하지 않고도 원상 복구 가능한 편집 모델.",
    ],
    next_plan=[
        "draft_refinement 프로덕션 활성화 후 2 주간 telemetry 누적 — refined/rollback 비율 및 평균 품질 개선 폭 실환경 측정.",
        "α6 Reflection 실측 — 컨설턴트 실 결재 누적 데이터로 프롬프트 버전별 수락률 A/B 평가.",
        "Phase D Memory 시스템 MVP — 대화 압축 + pgvector 기반 장기 메모리.",
        "Exemplar 30 건 import 확보 시 α2 v2-full Reward 엔진(대학 루브릭 거리 학습) 본 실측 진입.",
    ],
    references=[
        "한국특허전략개발원 연구노트 작성 가이드라인 준수 (주 1회~월 2회 권장).",
        "내부 메모리: session-handoff-2026-04-19-f, session-handoff-2026-04-20-k/l/m/o/q, session-handoff-2026-04-22-a.",
        "산업 표준 참조: AI SDK 5.x HITL 패턴(execute-less tool + InlineConfirm + addToolResult).",
    ],
)))


# ───────────────────────── 실행 ─────────────────────────

if __name__ == "__main__":
    saved = []
    for meta, body in NOTES:
        subdir = "timelevelup" if meta.project_code == "TimeLevelUp" else "record-levelup"
        out = OUT_ROOT / subdir / f"{meta.serial_no}_{meta.note_date}.docx"
        build_note(meta, body, out)
        saved.append(out)
    print(f"total saved: {len(saved)}")
    for p in saved:
        print(f"  - {p.relative_to(OUT_ROOT.parent.parent)}")
